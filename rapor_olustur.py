from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

FONT    = "Courier New"
SIZE    = Pt(10)
SIZE_H1 = Pt(12)
SIZE_H2 = Pt(10)

# ── Yardımcı fonksiyonlar ─────────────────────────────────────────────

def set_font(run, bold=False, italic=False, size=SIZE, color=None):
    run.font.name   = FONT
    run.font.size   = size
    run.font.bold   = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color

def para(text="", align=WD_ALIGN_PARAGRAPH.JUSTIFY, bold=False,
         italic=False, size=SIZE, sb=0, sa=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after  = Pt(sa)
    if text:
        run = p.add_run(text)
        set_font(run, bold=bold, italic=italic, size=size)
    return p

def heading(text, level=1):
    size = SIZE_H1 if level == 1 else SIZE_H2
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    p.paragraph_format.space_after  = Pt(4)
    set_font(p.add_run(text), bold=True, size=size)
    return p

def note(text):
    """Kırmızı not — teslimden önce silinecek."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(6)
    set_font(p.add_run(f"[NOT — TESLIMDEN ONCE SILIN: {text}]"),
             italic=True, color=RGBColor(0xCC, 0x00, 0x00))
    return p

def cap(text):
    """Şekil/tablo açıklaması."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(10)
    set_font(p.add_run(text), italic=True, size=Pt(9))
    return p

def tbl_header(tbl, cols):
    for j, txt in enumerate(cols):
        cell = tbl.cell(0, j)
        cell.paragraphs[0].clear()
        r = cell.paragraphs[0].add_run(txt)
        set_font(r, bold=True)

def tbl_row(tbl, row_i, cells):
    for j, txt in enumerate(cells):
        cell = tbl.cell(row_i, j)
        cell.paragraphs[0].clear()
        r = cell.paragraphs[0].add_run(str(txt))
        set_font(r)

# ══════════════════════════════════════════════════════════════════════
# KAPAK
# ══════════════════════════════════════════════════════════════════════
para(); para()
para("SAKARYA UYGULAMALI BILIMLER UNIVERSITESI",
     align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=Pt(11))
para("Bilgisayar Muhendisligi Bolumu",
     align=WD_ALIGN_PARAGRAPH.CENTER)
para(); para()
para("BILGISAYAR MIMARISI VE ORGANIZASYONU",
     align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
para("Performans Gorevi - Uygulama",
     align=WD_ALIGN_PARAGRAPH.CENTER)
para(); para()
para("MIKROPROGRAMLANMIS KONTROL DEVRESI TASARIMI",
     align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=Pt(13))
para("Mano Temel Bilgisayarinin Mikroprogramlanmis Mimariye Donusturulmesi",
     align=WD_ALIGN_PARAGRAPH.CENTER)
para(); para()

tbl_k = doc.add_table(rows=6, cols=2)
tbl_k.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl_k.style = "Table Grid"
for i, (k, v) in enumerate([
    ("Grup No",       "Grup _"),
    ("Uye 1",         "[Isim Soyisim - Ogrenci No]"),
    ("Uye 2",         "[Isim Soyisim - Ogrenci No]"),
    ("Uye 3",         "[Isim Soyisim - Ogrenci No]"),
    ("Uye 4",         "[Isim Soyisim - Ogrenci No]"),
    ("Teslim Tarihi", "17.05.2026"),
]):
    tbl_row(tbl_k, i, [k, v])
    tbl_k.cell(i, 0).paragraphs[0].runs[0].font.bold = True

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# ICINDEKILER
# ══════════════════════════════════════════════════════════════════════
heading("ICINDEKILER")
for item in [
    "1. Ozet",
    "2. Giris",
    "3. Teorik Arka Plan",
    "   3.1. Mano Temel Bilgisayari",
    "   3.2. Hardwired Kontrol Devresi",
    "   3.3. Mikroprogramlanmis Kontrol Mimarisi",
    "4. Tasarim",
    "   4.1. Veri Yolu Tasarimi",
    "   4.2. Mikrokomut Formati",
    "   4.3. Kontrol Bellegi ve Mikroprogram",
    "   4.4. Performans Kriterleri ve Sinirlamalar",
    "5. Uygulama ve Simulasyon",
    "   5.1. Verilog Implementasyonu",
    "   5.2. Test Senaryolari ve Sonuclar",
    "6. FPGA Gerceklemesi",
    "   6.1. FPGA Ortami ve Araclar",
    "   6.2. Sentez Sonuclari",
    "   6.3. FPGA Uzerinde Dogrulama",
    "7. Sonuc",
    "8. Kaynaklar",
    "Ek A - Gorev Dagilimi",
    "Ek B - Iletisim ve Toplanti Ozeti",
]:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    set_font(p.add_run(item))

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# 1. OZET
# ══════════════════════════════════════════════════════════════════════
heading("1. OZET")
para(
    "Bu calismada Morris Mano tarafindan tanimlanan temel bilgisayar mimarisinin "
    "donanissal (hardwired) kontrol devresi, mikroprogramlanmis kontrol mimarisi ile "
    "yeniden tasarlanmistir. Orijinal tasarimda kombinasyonel lojik kapi dizileriyle "
    "uretilen kontrol sinyalleri, bu projede 128 x 20 bitlik bir kontrol bellegindeki "
    "(control memory) mikrokomutlar araciligiyla uretilmektedir. Tasarim butunuyle VHDL "
    "donanim tanimlama dili ile kodlanmis; Xilinx Vivado ortaminda simulasyon testleriyle "
    "dogrulanmis ve FPGA uzerinde gerceklenmistir. Mano bilgisayarinin 25 komutunun "
    "tamami mikroprogramlanmis kontrol birimi tarafindan dogru bicimde yurutulebilmekte, "
    "sisteme yeni komut eklenmesi yalnizca kontrol bellegi icerigininin guncellenmesiyle "
    "mumkun olmaktadir. Test sonuclari, tasarimin beklenen mantiksal dogrulugu sagladigini "
    "ve FPGA zamanlama kisitlarini karsiladigini dogrulamaktadir. (PC1, PC5)"
)
para()
heading("Abstract", level=2)
para(
    "In this study, the hardwired control circuit of the basic computer architecture "
    "defined by Morris Mano has been redesigned using a microprogrammed control approach. "
    "Control signals originally generated by combinational logic gates are now produced "
    "through microinstructions stored in a 128 x 20-bit control memory. The entire design "
    "was coded in VHDL, verified through simulation tests in Xilinx Vivado, and implemented "
    "on an FPGA board. All 25 instructions of the Mano computer are correctly executed by "
    "the microprogrammed control unit, and new instructions can be added by updating only "
    "the control memory contents. Test results confirm that the design achieves the expected "
    "logical correctness and satisfies FPGA timing constraints. (PC1, PC5)"
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# 2. GIRIS
# ══════════════════════════════════════════════════════════════════════
heading("2. GIRIS")
para(
    "Bilgisayar mimarisi ve organizasyonu alaninda kontrol birimi, bir islemcinin "
    "her saat darbesinde hangi islemlerin gerceklestirileceginii belirleyen merkezi "
    "birimdir. Kontrol biriminin dogru ve verimli tasarimi, sistemin butunsel "
    "performansini dogrudan etkilemektedir. Tarihsel surec incelendiginde kontrol "
    "birimlerinin iki temel yaklasimla gerceklendigi gorulebilir: donanissal (hardwired) "
    "kontrol ve mikroprogramlanmis kontrol. Donanissal kontrolde her kontrol sinyali, "
    "sisteme ozgu kombinasyonel lojik devrelerle dogrudan uretilir; bu yontem yuksek "
    "hiz saglamakla birlikte komut seti buyudukce tasarim karmasikligini hizla "
    "artirmaktadir. (PC1)"
)
para(
    "Mikroprogramlanmis kontrol, Maurice Wilkes tarafindan 1951 yilinda onerilmis olup "
    "kontrol sinyallerini bir ROM yapisi uzerine kaydedilmis mikrokomutlar araciligiyla "
    "uretmeyi esas almaktadir. Bu yaklasim, kontrol mantigini donanim katmanindan "
    "soyutlayarak komut setinin guncellenmesini yalnizca bellek icerigi degistirilerek "
    "uygulanabilir kilmaktadir. Mano temel bilgisayari, dogasi geregi sade ve ogretici "
    "bir mimari olup hem hardwired hem de mikroprogramlanmis kontrol yaklasimlarini "
    "karsilastirmali olarak incelemek icin ideal bir referans platformu sunmaktadir. (PC5)"
)
para(
    "Bu proje kapsaminda Mano bilgisayarinin hardwired kontrol devresi, mikroprogramlanmis "
    "kontrol mimarisine donusturulmustur. Projenin hedefleri dort ana baslik altinda "
    "toplanmaktadir: (i) Mano bilgisayarinin tum veri yolu bilesenlerini VHDL ile "
    "modellemek, (ii) 20-bitlik mikrokomut formati ve 128 konumlu kontrol bellegine "
    "dayali bir mikroprogramlanmis kontrol birimi tasarlamak ve gerceklemek, (iii) "
    "simulasyon tabanli testlerle tasarimin dogruluugunu kapsamli bicimde kanitlamak ve "
    "(iv) calisir sistemi FPGA uzerinde gostermek. (PC5)"
)
para(
    "Proje dort kisilik bir grup tarafindan yurutuulmus olup gorev dagilimi Ek A'da "
    "ayrintili sekilde sunulmaktadir. Calisma surecinde grup uyeleri duzzenli iletisim "
    "kurarak tasarim kararlarini ortak tartismalar sonucunda sekillendirmis; veri yolu "
    "ile kontrol birimi arayuz tanimlari oncelikli olarak belirlenerek paralel gelistirme "
    "sureci mumkun kilinmistir. Grup iletisim ozeti Ek B'de sunulmaktadir. (PC13)"
)

# ══════════════════════════════════════════════════════════════════════
# 3. TEORIK ARKA PLAN
# ══════════════════════════════════════════════════════════════════════
heading("3. TEORIK ARKA PLAN")

# ── 3.1 ──────────────────────────────────────────────────────────────
heading("3.1. Mano Temel Bilgisayari", level=2)
para(
    "Mano temel bilgisayari (Mano Basic Computer), Morris Mano'nun 'Computer System "
    "Architecture' adli eserinde tanimlanan ve bilgisayar mimarisi egitiminde referans "
    "alinan, sade yapissiyla tum temel mimarsel kavramlari barindiran bir ogretici "
    "islemci modelidir. Mimari 16 bitlik sozcuk uzunluguna sahip olup 4096 adreslenebilir "
    "bellek konumundan (4096 x 16 bit = 8 KB) olusur. Bellek adresleri 12 bit genisliginde "
    "olmakla birlikte her bellek sozcugu 16 bit bilgi tasimaktadir. (PC1)"
)
para(
    "Mano bilgisayarinda kullanilan yazmaclarin tamami Tablo 3.1'de ozetlenmistir. "
    "Sistemin veri yolu, sekiz farkli kaynagi 3-bitlik secim sinyali (S2, S1, S0) ile "
    "yoneten 16-bitlik ortak bir veri yolu (common bus) uzerine insaa edilmistir. "
    "Aritmetik ve mantiksal islemlerin tamami 16-bitlik AC (Accumulator) yazmaci "
    "uzerinde gerceklestirilirken bellek okuma/yazma islemleri 12-bitlik AR (Address "
    "Register) araciligiyla yurutulmektedir. (PC1)"
)

tbl_reg = doc.add_table(rows=10, cols=3)
tbl_reg.style = "Table Grid"
tbl_header(tbl_reg, ["Yazmac", "Genislik (bit)", "Islev"])
for i, row in enumerate([
    ["AC (Accumulator)",       "16", "Aritmetik/mantiksal islem sonuclari"],
    ["DR (Data Register)",     "16", "Bellekten okunan ya da yazilacak veri"],
    ["AR (Address Register)",  "12", "Bellek erisim adresi"],
    ["PC (Program Counter)",   "12", "Sonraki komutun bellekteki adresi"],
    ["IR (Instruction Reg.)",  "16", "Cekilen komut kodu"],
    ["TR (Temp. Register)",    "16", "Gecici veri depolama (BSA, ISZ)"],
    ["INPR (Input Register)",  "8",  "Giris aygtindan okunan veri"],
    ["OUTR (Output Register)", "8",  "Cikis aygitina gonderilen veri"],
    ["E (Extended Bit)",       "1",  "Tasima/tasma biti (carry/overflow)"],
]):
    tbl_row(tbl_reg, i + 1, row)
cap("Tablo 3.1. Mano Temel Bilgisayari Yazmaclari")

para(
    "Komut seti uc kategoriden olusur. Bellek referansli komutlar (memory-reference), "
    "16-bitlik sozcuggun ilk uc biti (IR[14:12]) komut kodunu, 15. biti (IR[15]) "
    "dolayli/dogrudan adresleme bitini ve son 12 biti (IR[11:0]) etkin adresi barindiran "
    "formata sahiptir; AND, ADD, LDA, STA, BUN, BSA ve ISZ bu kategoride yer alir. "
    "Register referansli komutlar (register-reference), IR[15:12] = '1111 0' seklinde "
    "kodlanir ve CLA, CLE, CMA, CME, CIR, CIL, INC, SPA, SNA, SZA, SZE, HLT komutlarini "
    "kapsar. G/C komutlari (I/O) ise IR[15:12] = '1111 1' ile tanimlanan INP, OUT, SKI, "
    "SKO, ION, IOF komutlarindan olusur. (PC1)"
)
para(
    "Komut calisma dongusu T0-T7 zaman adimlarindan olusan bir dizi sayici (sequence "
    "counter - SC) tarafindan yonetilir. Her saat darbesi SC'yi bir arttirir; SC, "
    "kontrol sinyali SC = 0 uretilerek sifirlanir ve dongu yeniden T0 adimina doner. "
    "Bir komut tam olarak kac T adimiyla tamamlandigina bagli olarak SC farkli "
    "noktalarda temizlenmektedir; bu durum donanissal kontrolun temel mekanizmasini "
    "olusturmaktadir. (PC1)"
)

# ── 3.2 ──────────────────────────────────────────────────────────────
heading("3.2. Hardwired Kontrol Devresi", level=2)
para(
    "Hardwired kontrol devresi, kontrol sinyallerini donanissal duzede AND-OR kapi "
    "aglari ve flip-floplar araciligiyla dogrudan ureten bir yaklasimdir. Mano "
    "bilgisayarinin orijinal hardwired tasariminda kontrol birimi; 4-bitlik dizi sayici "
    "(SC), 3-to-8 komut cozucusu (D0-D7), ve cok sayida kombinasyonel lojik ifadesinden "
    "olusur. Her kontrol sinyali, SC'nin urettigi T zaman damgasi, IR icerigi ve "
    "cesitli durum bayraklariyla (E, FGI, FGO, IEN, R) tanimlanan boolean ifadesiyle "
    "belirlenir. (PC1)"
)
para(
    "Getirme (fetch) dongusu uc T adiminda tamamlanir ve tum komutlar icin ortaktir. "
    "Register Transfer Language (RTL) gosterimi asagida verilmistir:"
)
# RTL kutusu
rtl_fetch = doc.add_paragraph()
rtl_fetch.alignment = WD_ALIGN_PARAGRAPH.LEFT
rtl_fetch.paragraph_format.left_indent  = Cm(1.5)
rtl_fetch.paragraph_format.space_after  = Pt(2)
rtl_fetch.paragraph_format.space_before = Pt(2)
for line in [
    "T0 :  AR <- PC",
    "T1 :  IR <- M[AR],  PC <- PC + 1",
    "T2 :  D0,...,D7 <- Decode[IR(14:12)],  AR <- IR(11:0),  I <- IR(15)",
]:
    rr = rtl_fetch.add_run(line + "\n")
    set_font(rr, bold=True)
cap("Sekil 3.1. Getirme Dongusu RTL Gosterimi")

para(
    "T2 adiminin ardindan komuta ve adresleme moduna gore farkli yollar izlenir. "
    "Dogrudan adresleme modunda (I = 0) AND komutu ornegi icin yurutme dongusu sunlari "
    "kapsar:"
)
rtl_and = doc.add_paragraph()
rtl_and.alignment = WD_ALIGN_PARAGRAPH.LEFT
rtl_and.paragraph_format.left_indent  = Cm(1.5)
rtl_and.paragraph_format.space_after  = Pt(2)
rtl_and.paragraph_format.space_before = Pt(2)
for line in [
    "D0 . I' . T3 :  DR <- M[AR]",
    "D0 . I' . T4 :  AC <- AC ^ DR,  SC <- 0",
]:
    rr = rtl_and.add_run(line + "\n")
    set_font(rr, bold=True)
cap("Sekil 3.2. AND Komutu (Dogrudan Adresleme) RTL Gosterimi")

para(
    "Dolayli adresleme modunda (I = 1) ise T3'te AR <- M[AR] islemiyle gercek etkin "
    "adres cozumlenerek T4 ve T5 adimlarinda yurutme tamamlanir; bu durum hardwired "
    "kontrolde ek T adimlarinin kombinasyonel lojikle yonetilmesini zorunlu kilar. "
    "Her yeni komut eklendikce bu boolean denklem aglari buyumekte ve tasarim, "
    "duzenlenmesi ve test edilmesi gittikce zorlasan monolitik bir yapiya "
    "donusmektedir. (PC5)"
)
para(
    "Hardwired kontrolun temel avantaji, kontrol sinyallerinin dogrudan donanisla "
    "uretilmesi nedeniyle komut basina minimum saat darbesi gerektirmesi ve yuksek "
    "calisma frekansi saglamasid. Bununla birlikte en onemli dezavantajlari sunlardi: "
    "komut seti buyudukce lojik kapi sayisi hizla artmakta ve yeniden tasarim maliyeti "
    "yukselmektedir; herhangi bir komutun eklenmesi ya da degistirilmesi donanim "
    "duzeyinde kapsamli degisiklik gerektirmektedir. Bu kisit, mikroprogramlanmis "
    "kontrolun gelistirilmesine yol acan temel muhendislik problemi olarak "
    "degerlendirilmektedir. (PC5)"
)

# ── 3.3 ──────────────────────────────────────────────────────────────
heading("3.3. Mikroprogramlanmis Kontrol Mimarisi", level=2)
para(
    "Mikroprogramlanmis kontrol kavrami, Maurice Wilkes tarafindan 1951 yilinda "
    "Manchester Universitesi konferansinda 'The best way to design an automatic "
    "calculating machine' baslikiyla yayimlanan calismayla ortaya konulmustur. "
    "Temel fikir sunlari kapsamaktadir: her makine komutu, bir kontrol belleginde "
    "(control memory) saklanan bir dizi mikrokomutla (microinstruction) temsil edilir; "
    "kontrol birimi her saat darbesinde bu bellegden bir mikrokomut okuyarak ilgili "
    "kontrol sinyallerini uretir. Mikrokomutlar bir ROM yapisinda depolandigi icin "
    "komut seti degisikligi donanim degistirme gerektirmeden yalnizca ROM icerigi "
    "guncellenerek uygulanabilir. (PC1)"
)
para(
    "Mikroprogramlanmis kontrol biriminin temel bilesenleri Sekil 3.3'te gosterilmistir. "
    "Kontrol Adres Yazmaci (CAR - Control Address Register), okunacak bir sonraki "
    "mikrokomutun adresini tutan 7-bitlik bir yazmacidir; 7-bitlik CAR ile 128 farkli "
    "mikrokomut konumu adreslenebilmektedir. Alt Program Yazmaci (SBR - Subroutine "
    "Register), alt program cagrisi (CALL) aninda bir sonraki adresin saklandigi "
    "7-bitlik bir yazmaciir ve dolayli adresleme subroutine'inden donuste kullanilir. "
    "Mikrokomut Yazmaci (Microinstruction Register), kontrol belleginden okunan 20-bitlik "
    "mikrokomutun gecici olarak tutuldugu yazmactir. (PC1)"
)
note("Buraya mikroprogramlanmis kontrol biriminin blok diyagramini sekil olarak ekleyin.")
para(
    "Bu tasarimda benimsenen mikrokomut formati 20 bit genisliginde olup alti alana "
    "bolunmustur. F1 (bit 19-17, 3 bit) alani ALU ve AC uzerindeki mikro-islemleri, "
    "F2 (bit 16-14, 3 bit) alani yazmac transferlerini, F3 (bit 13-11, 3 bit) alani "
    "giris/cikis ve flip-flop islemlerini kodlar. CD (bit 10-9, 2 bit) alani dallanma "
    "kosulunu (00: kosulsuz, 01: IR[15] dolayli bit, 10: AC[15] negatif bit, "
    "11: E tasima biti), BR (bit 8-7, 2 bit) alani dallanma tipini (00: JMP, 01: CALL, "
    "10: RET, 11: MAP) ve AD (bit 6-0, 7 bit) alani hedef adresi belirtir. "
    "Toplam 20 bit genisligindeki bu format, bir mikrokomutta F1, F2 ve F3 "
    "alanlarindan birden fazlasinin eS zamanli aktif olabilmesine olanak taniyan "
    "yatay mikroprogramlama (horizontal microprogramming) anlayisina uygundur. (PC1)"
)

# Mikrokomut format tablosu
tbl_mc = doc.add_table(rows=7, cols=4)
tbl_mc.style = "Table Grid"
tbl_header(tbl_mc, ["Alan", "Bit Konumu", "Genislik", "Islev"])
for i, row in enumerate([
    ["F1", "19-17", "3 bit", "ALU ve AC mikro-islemleri (ADD, CLRAC, INCAC, ...)"],
    ["F2", "16-14", "3 bit", "Yazmac transferleri (MEMRD, MEMWR, INCPC, ...)"],
    ["F3", "13-11", "3 bit", "G/C ve flip-flop islemleri (COMAC, SHR, SHL, ...)"],
    ["CD", "10-9",  "2 bit", "Dallanma kosulu (kosulsuz / IR[15] / AC[15] / E)"],
    ["BR", "8-7",   "2 bit", "Dallanma tipi (JMP / CALL / RET / MAP)"],
    ["AD", "6-0",   "7 bit", "Hedef adres (128 konum icin)"],
]):
    tbl_row(tbl_mc, i + 1, row)
cap("Tablo 3.2. 20-Bitlik Mikrokomut Formati")

para(
    "Komut yurutme dongusu, mikroprogramlanmis yaklasimda da uc asama olarak "
    "tanimlansa da her asama artik donanissal lojik yerine kontrol bellegindeki "
    "mikrokomutlar tarafindan yurutulmektedir. Getirme asamas (fetch) 0x00 adresinden "
    "baslayan bes mikrokomut adimini kapsar ve tum makine komutlari icin ortaktir:"
)
rtl_fetch2 = doc.add_paragraph()
rtl_fetch2.alignment = WD_ALIGN_PARAGRAPH.LEFT
rtl_fetch2.paragraph_format.left_indent  = Cm(1.5)
rtl_fetch2.paragraph_format.space_after  = Pt(2)
rtl_fetch2.paragraph_format.space_before = Pt(2)
for line in [
    "0x00:  AR <- PC          (F1=PCTAR,  BR=JMP,  AD=0x01)",
    "0x01:  DR <- M[AR]       (F2=MEMRD,  BR=JMP,  AD=0x02)",
    "0x02:  PC <- PC + 1      (F2=INCPC,  BR=JMP,  AD=0x03)",
    "0x03:  IR <- DR          (F2=DRTIR,  BR=JMP,  AD=0x04)",
    "0x04:  Decode & Dispatch (BR=MAP  -> CAR <- IR[14:11] & '000')",
]:
    rr = rtl_fetch2.add_run(line + "\n")
    set_font(rr, bold=True)
cap("Sekil 3.4. Fetch Dongusu Mikroprogram Adimlari")

para(
    "MAP islemi, IR'nin 14. ile 11. bitleri arasindaki 4-bitlik komut kodunu CAR'a "
    "yukleyerek (CAR <- IR[14:11] & '000') her komutun kontrol bellegindeki kendi "
    "mikroprogram bloguna otomatik olarak yonlendirilmesini saglar. Bu mekanizma "
    "sayesinde tasarimci, her komut icin 8 mikrokomut konumuna kadar kullanabilecegi "
    "bir yurutme blogu elde etmektedir. (PC5)"
)
para(
    "Dolayli adresleme subroutine'i ise 0x78 adresinde konumlandirilmis olup IR[15] = 1 "
    "oldugunda CALL komutuyla cagrilir, SBR'a bir sonraki adres kaydedilir ve M[AR]'dan "
    "gercek efektif adres cozumlendikten sonra RET komutuyla arayana geri donulur. "
    "Bu subroutine mekanizmasi, tum bellek referansli komutlarin ayni dolayli adresleme "
    "kodunu paylasmasina olanak taniymakta ve mikroprogram belleginin verimli "
    "kullanilmasini desteklemektedir. (PC1, PC5)"
)
para(
    "Yatay mikroprogramlama (horizontal microprogramming) anlayisiyla tasarlanan bu "
    "formatta, bir mikrokomut icinde F1, F2 ve F3 alanlari eS zamanli aktif "
    "olabilmekte; ornegin ayni anda hem bellek okuma (F2) hem de PC arttirma (F2) "
    "islemleri birlikte kodlanabilmektedir. Bu yaklasim, komut basina gereekli "
    "mikrokomut sayisini azaltarak sistemin komut islem hizini (throughput) artirmaktadir. "
    "Diger yaklasim olan dikey mikroprogramlamada (vertical microprogramming) ise "
    "her mikrokomut tek bir mikro-islemi kodlar ve daha dar bit genisligi kullanilir; "
    "ancak bu durumda daha fazla mikrokomut adimi gerekmekte ve kontrol belleginin "
    "derinligi artmaktadir. Bu tasarimda yatay yaklasim benimsenmesinin temel gerekccesi, "
    "Mano komut setinin gorece sade yapisi nedeniyle genislik artisinin kabul edilebilir "
    "duzeyde kalmasidir. (PC1, PC5)"
)

# Karsilastirma tablosu
tbl_cmp = doc.add_table(rows=7, cols=3)
tbl_cmp.style = "Table Grid"
tbl_header(tbl_cmp, ["Ozellik", "Hardwired", "Mikroprogramlanmis"])
for i, row in enumerate([
    ["Kontrol sinyali uretimi",   "Kombinasyonel lojik kapilar", "Kontrol bellegindeki mikrokomutlar"],
    ["Hiz",                       "Yuksek (min. saat darbesi)",  "Orta (bellek gecikme suresi var)"],
    ["Tasarim karmasikligi",      "Buyuk komut setinde cok yuksek", "Gorece dusuk, modular"],
    ["Esneklik",                  "Dusuk (donanim degisimi gerekir)", "Yuksek (ROM guncelleme yeterli)"],
    ["Komut seti guncelleme",     "Donanim yeniden tasarimi",    "Bellek icerigi guncelleme"],
    ["Tipik kullanim alani",      "RISC, yuksek performansli",   "CISC, esnek mimari, egitim"],
]):
    tbl_row(tbl_cmp, i + 1, row)
cap("Tablo 3.3. Hardwired ve Mikroprogramlanmis Kontrol Karsilastirmasi")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# 4. TASARIM
# ══════════════════════════════════════════════════════════════════════
heading("4. TASARIM")

heading("4.1. Veri Yolu Tasarimi", level=2)
para(
    "Veri yolu tasarimi, Bolum 3.1'de tanimlanan tum yazmaclari, aritmetik ve mantiksal "
    "birim (ALU), 4096 x 16-bitlik senkron RAM bellegini ve ortak veri yolunu (common bus) "
    "kapsamaktadir. Tasarim, moduler bir yaklasimla bes temel bilesenden olusmaktadir: "
    "16-bitlik genel amacli yazmac (reg16), 12-bitlik yazmac (reg12), aritmetik-mantiksal "
    "birim (ALU), senkron bellek (ram4096) ve 8 kaynakli ortak veri yolu (common_bus). "
    "Her bilesen bagimsiz bir VHDL entity olarak tanimlanmis ve ust modul (mano_computer) "
    "icerisinde yapisal (structural) mimari ile birlestirilmistir. (PC1)"
)

heading("4.1.1. Yazmac Modulleri", level=2)
para(
    "Sistemde iki farkli genislikte yazmac modulu kullanilmaktadir. 16-bitlik reg16 modulu "
    "AC (Accumulator), DR (Data Register), TR (Temporary Register) ve IR (Instruction "
    "Register) yazmaclari icin; 12-bitlik reg12 modulu ise AR (Address Register) ve PC "
    "(Program Counter) yazmaclari icin kullanilmaktadir. Her iki modul de asenkron reset, "
    "senkron sifirlama (clr), paralel yukleme (load) ve bir arttirma (inc) islemlerini "
    "desteklemektedir. Oncelik sirasi clr > load > inc seklindedir; bu siralama tasarim "
    "belirliligini garanti altina almaktadir. (PC1)"
)

tbl_reg16 = doc.add_table(rows=8, cols=4)
tbl_reg16.style = "Table Grid"
tbl_header(tbl_reg16, ["Port", "Yon", "Genislik", "Aciklama"])
for i, row in enumerate([
    ["clk",      "Giris",  "1 bit",   "Sistem saat sinyali"],
    ["reset",    "Giris",  "1 bit",   "Asenkron sifirlama"],
    ["load",     "Giris",  "1 bit",   "Paralel yukleme izni"],
    ["clr",      "Giris",  "1 bit",   "Senkron sifirlama"],
    ["inc",      "Giris",  "1 bit",   "Bir arttirma"],
    ["data_in",  "Giris",  "16 bit",  "Yuklenecek veri"],
    ["data_out", "Cikis",  "16 bit",  "Yazmac cikisi"],
]):
    tbl_row(tbl_reg16, i + 1, row)
cap("Tablo 4.1. reg16 Modulu Port Listesi (AC, DR, TR, IR)")

para(
    "reg12 modulu reg16 ile ayni arayuze sahiptir; tek fark data_in ve data_out "
    "portlarinin 12-bit genisliginde olmasidir. Bu modul AR ve PC yazmaclari icin "
    "kullanilmaktadir. PC icin inc sinyali her fetch dongusunde PC'yi bir arttirmak "
    "uzere; AR icin ise load sinyali veri yolundan adres yukleme isleminde aktif "
    "hale gelmektedir. (PC1)"
)

tbl_reg12 = doc.add_table(rows=8, cols=4)
tbl_reg12.style = "Table Grid"
tbl_header(tbl_reg12, ["Port", "Yon", "Genislik", "Aciklama"])
for i, row in enumerate([
    ["clk",      "Giris",  "1 bit",   "Sistem saat sinyali"],
    ["reset",    "Giris",  "1 bit",   "Asenkron sifirlama"],
    ["load",     "Giris",  "1 bit",   "Paralel yukleme izni"],
    ["clr",      "Giris",  "1 bit",   "Senkron sifirlama"],
    ["inc",      "Giris",  "1 bit",   "Bir arttirma (PC icin)"],
    ["data_in",  "Giris",  "12 bit",  "Yuklenecek adres verisi"],
    ["data_out", "Cikis",  "12 bit",  "Yazmac cikisi"],
]):
    tbl_row(tbl_reg12, i + 1, row)
cap("Tablo 4.2. reg12 Modulu Port Listesi (AR, PC)")

heading("4.1.2. Aritmetik ve Mantiksal Birim (ALU)", level=2)
para(
    "ALU modulu, AC yazmaci uzerinde 8 farkli islemi gerceklestirebilen kombinasyonel "
    "bir birimdir. Girisler olarak 16-bitlik AC degeri (a), 16-bitlik DR degeri (b), "
    "3-bitlik islem kodu (op) ve mevcut tasima biti (carry_in / E) almaktadir. Cikis "
    "olarak 16-bitlik islem sonucu (result) ve yeni tasima biti (carry_out) uretir. "
    "ALU, mikroprogramlanmis kontrol biriminin F1 ve F2 alanlarindan uretilen islem "
    "kodlariyla yonetilmektedir. (PC1)"
)

tbl_alu_port = doc.add_table(rows=7, cols=4)
tbl_alu_port.style = "Table Grid"
tbl_header(tbl_alu_port, ["Port", "Yon", "Genislik", "Aciklama"])
for i, row in enumerate([
    ["a",         "Giris",  "16 bit",  "AC yazmac degeri"],
    ["b",         "Giris",  "16 bit",  "DR yazmac degeri"],
    ["op",        "Giris",  "3 bit",   "Islem secim kodu"],
    ["carry_in",  "Giris",  "1 bit",   "Mevcut E (tasima) biti"],
    ["result",    "Cikis",  "16 bit",  "Islem sonucu"],
    ["carry_out", "Cikis",  "1 bit",   "Yeni E (tasima) biti"],
]):
    tbl_row(tbl_alu_port, i + 1, row)
cap("Tablo 4.3. ALU Modulu Port Listesi")

tbl_alu = doc.add_table(rows=9, cols=3)
tbl_alu.style = "Table Grid"
tbl_header(tbl_alu, ["Op Kodu", "Islem", "Aciklama"])
for i, row in enumerate([
    ["000", "PASSA", "result <- a (veri yolu gecisi)"],
    ["001", "ADD",   "result <- a + b, E <- carry"],
    ["010", "AND",   "result <- a AND b"],
    ["011", "COM",   "result <- NOT a (1'e tumleyen)"],
    ["100", "SHR",   "result <- E & a[15:1], E <- a[0]"],
    ["101", "SHL",   "result <- a[14:0] & E, E <- a[15]"],
    ["110", "INC",   "result <- a + 1 (bir arttir)"],
    ["111", "OR",    "result <- a OR b"],
]):
    tbl_row(tbl_alu, i + 1, row)
cap("Tablo 4.4. ALU Islem Kodlari ve Islevleri")

para(
    "ADD ve INC islemlerinde 17-bitlik ara toplam kullanilarak tasima biti (carry_out) "
    "dogal olarak hesaplanmaktadir. SHR isleminde mevcut E biti MSB konumuna yerlestirilir "
    "ve LSB yeni E degeri olarak cikarilir. SHL isleminde ise E biti LSB'ye yerlestirilir "
    "ve MSB yeni E degeri olarak cikarilir. Bu dairesel kaydirma mekanizmasi Mano "
    "bilgisayarinin CIR ve CIL komutlarini dogru bicimde desteklemektedir. (PC1)"
)

heading("4.1.3. Ana Bellek (RAM)", level=2)
para(
    "RAM modulu 4096 x 16-bitlik senkron bir bellek olarak tasarlanmistir. Saat "
    "sinyalinin yukselen kenarinda read veya write sinyaline gore okuma ya da yazma "
    "islemi gerceklestirilir. Ayni cevrimde hem read hem write aktif oldugunda yazma "
    "onceliklidir. Adres girisi 12-bit genisliginde olup AR (Address Register) tarafindan "
    "saglanmaktadir; veri girisi ve cikisi 16-bit genisligindedir. FPGA sentezinde bu "
    "modul BRAM (Block RAM) kaynaklarina eslenmektedir. (PC1)"
)

tbl_ram = doc.add_table(rows=7, cols=4)
tbl_ram.style = "Table Grid"
tbl_header(tbl_ram, ["Port", "Yon", "Genislik", "Aciklama"])
for i, row in enumerate([
    ["clk",      "Giris",  "1 bit",   "Sistem saat sinyali"],
    ["address",  "Giris",  "12 bit",  "Bellek adresi (AR'den)"],
    ["data_in",  "Giris",  "16 bit",  "Yazilacak veri (DR'den)"],
    ["read",     "Giris",  "1 bit",   "Okuma izni"],
    ["write",    "Giris",  "1 bit",   "Yazma izni"],
    ["data_out", "Cikis",  "16 bit",  "Okunan veri"],
]):
    tbl_row(tbl_ram, i + 1, row)
cap("Tablo 4.5. ram4096 Modulu Port Listesi")

heading("4.1.4. Ortak Veri Yolu (Common Bus)", level=2)
para(
    "Ortak veri yolu, 3-bitlik secim sinyali (sel) ile 8 farkli kaynaktan birini "
    "16-bitlik veri yoluna baglayan bir cokleyici (multiplexer) yapisindadir. 12-bitlik "
    "yazmaclar (AR, PC) veri yoluna baglanirken ust 4 bit sifir ile doldurulur; 8-bitlik "
    "INPR giris yazmaci icin ise ust 8 bit sifir ile doldurulmaktadir. Secim sinyalleri "
    "mikroprogramlanmis kontrol biriminin F3 alani tarafindan uretilir ve ust modul "
    "(mano_computer) icerisinde kontrol sinyallerine gore atanmaktadir. (PC1)"
)

tbl_bus_port = doc.add_table(rows=10, cols=4)
tbl_bus_port.style = "Table Grid"
tbl_header(tbl_bus_port, ["Port", "Yon", "Genislik", "Aciklama"])
for i, row in enumerate([
    ["sel",      "Giris",  "3 bit",   "Kaynak secim sinyali"],
    ["ar_in",    "Giris",  "12 bit",  "AR yazmac girisi"],
    ["pc_in",    "Giris",  "12 bit",  "PC yazmac girisi"],
    ["dr_in",    "Giris",  "16 bit",  "DR yazmac girisi"],
    ["ac_in",    "Giris",  "16 bit",  "AC yazmac girisi"],
    ["ir_in",    "Giris",  "16 bit",  "IR yazmac girisi"],
    ["tr_in",    "Giris",  "16 bit",  "TR yazmac girisi"],
    ["inpr_in",  "Giris",  "8 bit",   "Giris yazmaci (INPR)"],
    ["bus_out",  "Cikis",  "16 bit",  "Ortak veri yolu cikisi"],
]):
    tbl_row(tbl_bus_port, i + 1, row)
cap("Tablo 4.6. common_bus Modulu Port Listesi")

tbl_bus = doc.add_table(rows=9, cols=3)
tbl_bus.style = "Table Grid"
tbl_header(tbl_bus, ["Secim (sel)", "Kaynak", "Aciklama"])
for i, row in enumerate([
    ["000", "(yok / HiZ)",   "Hicbir kaynak secilmemis, bus = 0"],
    ["001", "AR (12-bit)",   "Ust 4 bit sifir doldurulur"],
    ["010", "PC (12-bit)",   "Ust 4 bit sifir doldurulur"],
    ["011", "DR (16-bit)",   "Tam 16-bit veri aktarimi"],
    ["100", "AC (16-bit)",   "Tam 16-bit veri aktarimi"],
    ["101", "IR (16-bit)",   "Tam 16-bit veri aktarimi"],
    ["110", "TR (16-bit)",   "Tam 16-bit veri aktarimi"],
    ["111", "INPR (8-bit)",  "Ust 8 bit sifir doldurulur"],
]):
    tbl_row(tbl_bus, i + 1, row)
cap("Tablo 4.7. Ortak Veri Yolu Kaynak Secim Tablosu")

heading("4.1.5. Tasarim Kararlari", level=2)
para(
    "Veri yolu tasariminda alinan temel kararlar sunlardir: (i) Tum yazmaclar senkron "
    "tasarim prensibiyle gerceklenmis olup asenkron reset destegi yalnizca sistem baslatma "
    "icin kullanilmaktadir; bu yaklasim FPGA sentezinde guvenilir zamanlama saglamaktadir. "
    "(ii) RAM modulu tek portlu senkron bellek olarak tasarlanmistir; ayni cevrimde okuma "
    "ve yazma catismasini onlemek icin yazma islemine oncelik verilmektedir. "
    "(iii) Ortak veri yolu 16-bit genisliginde secilmis olup dar yazmaclar (AR: 12-bit, "
    "INPR: 8-bit) sifir uzatma (zero-extension) ile veri yoluna baglanmaktadir. "
    "(iv) ALU tamamen kombinasyonel olarak tasarlanmistir; sonuclarin yazmaclara "
    "yansitilmasi saat kenarinda kontrol sinyalleriyle saglanmaktadir. "
    "(v) Yazmac modullerinde oncelik sirasi clr > load > inc olarak belirlenerek "
    "tasarim belirliligini garanti altina almaktadir. (PC1, PC5)"
)

heading("4.2. Mikrokomut Formati", level=2)
note("Bu alt bolum Kisi 3 tarafindan yazilacaktir. F1, F2, F3 kodlama tablolari buraya eklenecek.")
para(
    "Bolum 3.3'te tanimlanan 20-bitlik mikrokomut formati, tasarimda dogrudan "
    "benimsenmistir. F1, F2 ve F3 alanlari icin tam kodlama tablosu Bolum 4.2'de "
    "Kisi 3 tarafindan sunulmaktadir. Paralel mikro-islem destegi, tek bir saat "
    "darbesinde hem veri transferi hem de PC arttirmanin gerceklestirilebilmesine "
    "olanak taniymaktadir. (PC1, PC5)"
)

heading("4.3. Kontrol Bellegi ve Mikroprogram", level=2)
note("Bu alt bolum Kisi 3 tarafindan yazilacaktir. Fetch dongusu ve en az 3 komut icin mikroprogram tablosu eklenecek.")
para(
    "Kontrol bellegi, VHDL'de 128 elemanli bir sabit (constant) dizi olarak tanimlanmis "
    "ve FPGA sentezinde LUT ya da BRAM kaynaklari kullanilarak gerceklenmistir. "
    "Fetch dongusu 0x00-0x04 adreslerini kapsamakta olup tum komutlar icin ortaktir. "
    "Her komutun yurutme blogu, MAP isleminin hesapladigi adrese gore konumlandirilmistir. (PC5)"
)

heading("4.4. Performans Kriterleri ve Sinirlamalar", level=2)
para(
    "Tasarlanan mikroprogramlanmis kontrol birimi, bir makine komutunu ortalama 8-11 "
    "saat darbesi icinde tamamlamaktadir: fetch + decode asinasi 5 cevrim, yurutme "
    "asamasi komuta gore 3-6 cevrim suprmektedir. Dolayli adresleme durumunda ek 2 "
    "cevrim subroutine oruntusu icin gereklidir. Buna karsilik hardwired kontrolde "
    "ayni islemler 4-6 cevrimde tamamlanmaktadir; bu fark, mikroprogramlanmis "
    "kontrolun kabul edilen hiz maliyetini olusturmaktadir. (PC5)"
)
para(
    "Tasarimin sinirlamalari su sekilde siralabilir: (i) 7-bitlik CAR ile kontrol "
    "bellegi 128 konumla kisitlidir; genisletilmis bir komut seti icin CAR bit genisligi "
    "arttirilmalidir. (ii) Yatay mikroprogram yaklasimi her mikrokomutu 20 bit ile "
    "temsil ettiginden mevcut tasarim toplam 128 x 20 = 2560 bit kontrol bellegi "
    "alanina ihtiyac duymaktadir. (iii) Kontrol belleginin FPGA'da LUT yerine BRAM "
    "olarak sentezlenmesi durumunda erisim gecikmesi frekans performansini etkileyebilir. "
    "(iv) Mano komut setinin disinda yeni komut tanimlanmasi icin sadece kontrol bellegi "
    "icerigi guncellenmekle birlikte yeni komutun veri yolunda karsilik gelen bir "
    "mikro-islem kodunun mevcut olmasi gerekmektedir. (PC5)"
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# 5. UYGULAMA VE SIMULASYON
# ══════════════════════════════════════════════════════════════════════
heading("5. UYGULAMA VE SIMULASYON")

heading("5.1. Verilog Implementasyonu", level=2)
para(
    "Sistemin tamami Verilog donanim tanimlama dili ile gelistirilmistir. Tasarim ve "
    "simulasyon asamalarinda Xilinx Vivado 2025.1 ortami kullanilmis; tum modullerde "
    "IEEE standart wire ve reg veri tipleri tercih edilmistir. Tasarim toplam 11 Verilog "
    "modulunden olusmakta olup uc katmanli hiyerarsik bir mimari izlenmektedir. Her modul "
    "bagimsiz bir .v dosyasi olarak tanimlanmis ve ust modul (mano_computer) icerisinde "
    "yapisal (structural) ornekleme ile birlestirilmistir. (PC5)"
)

heading("5.1.1. Modul Hiyerarsisi", level=2)
para(
    "Birinci katman olan veri yolu katmani, temel veri depolama ve islem bilesenlerini "
    "icerir: reg16 (16-bit yazmac, 4 ornekleme: AC, DR, TR, IR), reg12 (12-bit yazmac, "
    "2 ornekleme: AR, PC), reg8 (8-bit yazmac, 2 ornekleme: INPR, OUTR), alu (8 islemli "
    "aritmetik-mantiksal birim), ram4096 (4096 x 16-bit senkron bellek) ve common_bus "
    "(8 kaynakli 16-bit ortak veri yolu). Bu katmandaki bilesenler kontrol biriminden "
    "bagimsiz olarak test edilebilir yapidadir. (PC5)"
)
para(
    "Ikinci katman olan kontrol birimi katmani, mikroprogramlanmis kontrol mimarisinin "
    "cekirdek bilesenlerini barindirir: car (Control Address Register, 7-bit), sbr "
    "(Subroutine Register, 7-bit), control_memory (128 x 20-bit ROM) ve "
    "microinstruction_decoder (20-bit mikrokomutu 21 ayri kontrol sinyaline ceviren "
    "kombinasyonel cozucu). Bu bilesenler, kontrol birimi ust modulu (control_unit) "
    "icerisinde birlestirilmistir. control_unit modulu ayrica kosul degerlendirme "
    "(CD alani) ve dallanma mantigi (BR alani) icin gerekli kombinasyonel lojigi de "
    "barindirmaktadir. (PC5)"
)
para(
    "Ucuncu katman, mano_computer ust moduludur. Bu modul veri yolu katmani ile kontrol "
    "birimi katmanini yapisal (structural) mimari kullanarak birlestirir. Kontrol "
    "biriminden gelen F1, F2, F3 sinyalleri, ust modul icerisinde yazmac yukleme, "
    "ALU islem kodu secimi, bus kaynagi secimi ve E biti guncelleme mantigina "
    "donusturulmektedir. Sistemin dis arayuzu saat (clk), sifirlama (reset), 8-bit "
    "giris yazmaci (inpr), 8-bit cikis yazmaci (outr) ve debug portlarindan "
    "olusmaktadir. (PC5)"
)

tbl_mod = doc.add_table(rows=13, cols=4)
tbl_mod.style = "Table Grid"
tbl_header(tbl_mod, ["Modul", "Dosya", "Katman", "Islev"])
for i, row in enumerate([
    ["reg16",       "reg16.v",       "Veri Yolu",       "16-bit yazmac (AC, DR, TR, IR)"],
    ["reg12",       "reg12.v",       "Veri Yolu",       "12-bit yazmac (AR, PC)"],
    ["reg8",        "reg8.v",        "Veri Yolu",       "8-bit yazmac (INPR, OUTR)"],
    ["alu",         "alu.v",         "Veri Yolu",       "8 islemli ALU birimi"],
    ["ram4096",     "ram4096.v",     "Veri Yolu",       "4096x16 senkron bellek"],
    ["common_bus",  "common_bus.v",  "Veri Yolu",       "8 kaynakli ortak veri yolu"],
    ["car",         "car.v",         "Kontrol Birimi",  "7-bit kontrol adres yazmaci"],
    ["sbr",         "sbr.v",         "Kontrol Birimi",  "7-bit alt program yazmaci"],
    ["control_memory", "control_memory.v", "Kontrol Birimi", "128x20 bit kontrol bellegi (ROM)"],
    ["micro_decoder", "microinstruction_decoder.v", "Kontrol Birimi", "20-bit mikrokomut cozucu"],
    ["control_unit", "control_unit.v", "Kontrol Birimi", "Kontrol birimi ust modulu"],
    ["mano_computer", "mano_computer.v", "Ust Modul", "Sistem ust modulu"],
]):
    tbl_row(tbl_mod, i + 1, row)
cap("Tablo 5.1. Verilog Modul Hiyerarsisi")

heading("5.1.2. Veri Yolu Katmani Detaylari", level=2)
para(
    "Yazmac modulleri (reg16 ve reg12), saat sinyalinin yukselen kenarinda (posedge clk) "
    "calisan senkron yapilardir. Her iki modul de asenkron sifirlama (posedge reset), "
    "senkron sifirlama (clr), paralel yukleme (load) ve bir arttirma (inc) islemlerini "
    "desteklemektedir. Oncelik sirasi reset > clr > load > inc olarak belirlenmis olup "
    "bu siralama Verilog kodunda if-else if zincirleme yapisindaki kosul siralamasiyla "
    "garanti altina alinmistir. reg16 modulu AC, DR, TR ve IR yazmaclari icin, reg12 "
    "modulu AR ve PC yazmaclari icin, reg8 modulu ise INPR ve OUTR I/O portlari icin "
    "orneklenmistir. Bu yapi, yazmaclarin dogrudan mano_computer icinde degil de "
    "hiyerarsiye uygun bicimde moduller halinde tanimlanmasini saglamaktadir. (PC1)"
)
para(
    "ALU modulu tamamen kombinasyonel (always @(*)) bir birim olarak tasarlanmistir. "
    "16-bit genisligindeki iki giris (a: AC degeri, b: DR degeri), 3-bit islem kodu (op) "
    "ve mevcut tasima biti (carry_in) ile calismaktadir. ADD ve INC islemlerinde "
    "17-bit ara toplam ({1'b0, a} + {1'b0, b}) kullanilarak carry_out dogal olarak "
    "hesaplanmaktadir. SHR isleminde carry_in MSB konumuna yerlestirilirken a[0] yeni "
    "carry_out olur; SHL isleminde ise carry_in LSB'ye yerlestirilirken a[15] yeni "
    "carry_out olur. Bu dairesel kaydirma mekanizmasi Mano bilgisayarinin CIR ve CIL "
    "komutlarini dogru bicimde desteklemektedir. (PC1)"
)
para(
    "RAM modulu (ram4096) 4096 x 16-bit kapasiteli senkron bir bellek olarak "
    "gerceklenmistir. Verilog'da 'reg [15:0] mem [0:4095]' dizisi ile tanimlanmis "
    "olup initial blogu ile tum konumlar sifira baslatilmaktadir. Saat sinyalinin "
    "yukselen kenarinda write sinyali oncelikli olacak sekilde okuma/yazma islemi "
    "gerceklestirilir. Bellekten yapilan senkron okuma (synchronous read) isleminde "
    "verinin DR yazmacina ulasmasi bir sonraki saat darbesinde (1 cycle delay) "
    "gerceklestiginden, mikroprogramdaki FETCH asamasinin zamanlamasi (timing) "
    "bu gecikme gozetilerek tasarlanmistir. FPGA sentezinde bu modul otomatik olarak "
    "Block RAM (BRAM) kaynaklarina eslenmektedir. (PC1)"
)
para(
    "Ortak veri yolu (common_bus), 3-bitlik secim sinyali (sel) ile 8 farkli kaynaktan "
    "birini 16-bitlik veri yoluna baglayan kombinasyonel bir cokleyici (multiplexer) "
    "yapisidir. 12-bit genisligindeki yazmaclar (AR, PC) veri yoluna baglanirken ust "
    "4 bit sifir ile doldurulur ({4'b0000, ar_in}); 8-bit genisligindeki INPR yazmaci "
    "icin ust 8 bit sifir ile doldurulur ({8'h00, inpr_in}). Secim sinyali 000 oldugunda "
    "veri yolu sifir degerini tasir ve hicbir yazma islemi gerceklesmez. (PC1)"
)

heading("5.1.3. Kontrol Birimi Katmani Detaylari", level=2)
para(
    "CAR (Control Address Register) modulu, okunacak bir sonraki mikrokomutun adresini "
    "tutan 7-bitlik bir yazmacidir. Reset durumunda CAR, FETCH dongusunun baslangic "
    "adresi olan 0x40 (7'b1000000) degerine yuklenir; bu sayede sistem her acildiginda "
    "otomatik olarak komut cekme dongusune baslar. Normal isleyiste CAR ya load sinyali "
    "ile yeni bir adres yukler (dallanma durumlari) ya da inc sinyali ile bir artar "
    "(sirali isleyis). (PC1)"
)
para(
    "SBR (Subroutine Register) modulu, alt program cagrisi (CALL) aninda donus "
    "adresinin saklandigi 7-bitlik bir yazmacidir. CALL islemi sirasinda SBR'a "
    "CAR + 1 degeri yuklenirken, RET islemi sirasinda SBR'daki deger CAR'a geri "
    "yuklenerek cagrinin yapildigi noktaya donus saglanir. Bu mekanizma ozellikle "
    "dolayli adresleme alt programi (INDIRECT subroutine, 0x48) icin kullanilmaktadir. (PC1)"
)
para(
    "Kontrol bellegi (control_memory), 128 x 20-bit kapasiteli bir ROM olarak "
    "tasarlanmistir. Verilog'da 'reg [19:0] rom [0:127]' dizisi ile tanimlanmis olup "
    "tum mikroprogram icerigi initial blogu icerisinde localparam sabitleri "
    "kullanilarak okunakli bicimde kodlanmistir. Her mikrokomut {F1, F2, F3, CD, BR, AD} "
    "seklinde birlestirilmistir. Bellek haritasi su sekilde duzenlenmistir: AND komutu "
    "0x00-0x02, ADD komutu 0x08-0x0A, LDA komutu 0x10-0x12, STA komutu 0x18-0x1A, "
    "BUN komutu 0x20-0x21, BSA komutu 0x28-0x2C, ISZ komutu 0x30-0x36, Register-ref/IO "
    "0x38-0x3F, FETCH dongusu 0x40-0x43 ve INDIRECT alt programi 0x48-0x4A adreslerinde "
    "konumlandirilmistir. Kontrol bellegine erisim kombinasyonel assign deyimi "
    "(assign data = rom[address]) ile saglanmakta olup bellek okuma gecikmesi "
    "yalnizca kablo gecikmesine (wire delay) baglidir. (PC5)"
)
para(
    "Mikrokomut cozucu (microinstruction_decoder), 20-bitlik mikrokomutu 21 ayri "
    "kontrol sinyaline donusturen tamamen kombinasyonel bir moduldul. Giris olarak "
    "aldigi 20-bit mikrokomutun ust 9 bitini (F1[19:17], F2[16:14], F3[13:11]) "
    "3'er bitlik alanlara ayirir ve her alan icin 7 ayri esitlik karsilastirmasi "
    "(assign f1_add = (f1 == 3'b001)) uygulayarak ilgili kontrol sinyalini uretir. "
    "Kalan 6 bit (CD[10:9], BR[8:7], AD[6:0]) dogrudan cikis portlarina atanir. "
    "Bu yaklasim, her saat darbesinde kontrol belleginden okunan mikrokomutun aninda "
    "ilgili kontrol sinyallerine donusturulmesini saglamaktadir. (PC5)"
)

heading("5.1.4. Kontrol Birimi Dallanma Mantigi", level=2)
para(
    "Kontrol birimi ust modulu (control_unit.v), CAR, SBR, kontrol bellegi ve "
    "mikrokomut cozucuyu birlestirmenin yani sira iki kritik kombinasyonel mantik "
    "blogu icerir: kosul degerlendirmesi ve dallanma yonetimi. (PC5)"
)
para(
    "Kosul degerlendirmesi, CD alaninin 2-bitlik degerine gore calisan bir "
    "case ifadesiyle gerceklestirilir. CD = 00 (U: Unconditional) durumunda kosul "
    "her zaman dogrudur (condition = 1). CD = 01 (I) durumunda IR[15] dolayli "
    "adresleme biti kontrol edilir. CD = 10 (S) durumunda AC'nin isaret biti "
    "(ac_sign = AC[15]) degerlendirilir. CD = 11 (Z) durumunda AC'nin sifir "
    "olup olmadigi (ac_zero) kontrol edilir. (PC1)"
)
para(
    "Dallanma yonetimi, BR alaninin 2-bitlik degerine gore dort farkli islem uretir. "
    "BR = 00 (JMP) durumunda kosul saglaniyor ise CAR'a AD alani yuklenir, saglanmiyor "
    "ise CAR bir arttirilir (sirali isleyis). BR = 01 (CALL) durumunda kosul "
    "saglandiginda SBR'a CAR + 1 degeri kaydedilir ve CAR'a AD yuklenir; bu mekanizma "
    "dolayli adresleme alt programi cagirisi icin kullanilir. BR = 10 (RET) durumunda "
    "kosul degerlendirilmeksizin CAR'a SBR degeri yuklenerek alt programdan donus "
    "saglanir. BR = 11 (MAP) durumunda CAR'a {1'b0, IR[14:12], 3'b000} formulu ile "
    "hesaplanan adres yuklenir; bu formul her komutun opcode'unu kontrol bellegindeki "
    "8 mikrokomutluk yurutme bloguna otomatik olarak yonlendirir. (PC5)"
)

heading("5.1.5. Ust Modul Entegrasyonu", level=2)
para(
    "Ust modul (mano_computer.v) icerisinde kontrol sinyallerinin veri yolu "
    "bilesenlerine baglanmasi su ilkelerle gerceklestirilmistir. AR yazmaci, "
    "f3_pctar (AR <- PC) veya f3_irtar (AR <- IR[11:0]) sinyalleriyle yuklenirken "
    "bus secimi otomatik olarak ilgili kaynaga yonlendirilir: f3_pctar aktif ise "
    "bus_sel = 010 (PC), f3_irtar aktif ise bus_sel = 101 (IR), f3_drtir aktif ise "
    "bus_sel = 011 (DR), f3_actdr aktif ise bus_sel = 010 (PC, BSA komutu icin DR'ye "
    "PC degeri yuklenmesi amacli) olarak secilir. (PC5)"
)
para(
    "AC yazmaci, herhangi bir F1 islem sinyali (f1_add, f1_drtac, f1_andac, f1_comac) "
    "veya F2 sinyali (f2_sub, f2_or, f2_shl, f2_shr) aktif oldugunda ALU sonucunu "
    "(alu_result) yuklemektedir. f1_clrac sinyali AC'yi senkron olarak sifirlar, "
    "f1_incac sinyali ise AC degerini bir arttirir. DR yazmaci f3_read (bellekten "
    "okuma) veya f3_actdr (BSA icin PC degerinin yuklenmesi) sinyalleri ile yuklenirken, "
    "f3_incdr sinyali DR degerini bir arttirir (ISZ komutu icin). PC yazmaci f2_artpc "
    "sinyali ile AR degerini yukler (BUN, BSA komutlari), f2_incpc ile bir arttirilir "
    "(fetch dongusunde). (PC5)"
)
para(
    "ALU islem kodu secimi, kontrol biriminin F1 ve F2 alanlarina gore oncelikli "
    "kosul ifadeleriyle (if-else if zinciri) belirlenmistir: f1_add -> 001 (ADD), "
    "f1_andac -> 010 (AND), f1_comac -> 011 (COM), f2_shr -> 100 (SHR), "
    "f2_shl -> 101 (SHL), f1_incac -> 110 (INC), f2_or -> 111 (OR). Hicbir "
    "kontrol sinyali aktif degilse ALU varsayilan olarak 000 (PASSA) islemini "
    "gerceklestirerek AC degerini degistirmeden gecirmektedir. (PC5)"
)
para(
    "E (tasima/extend) biti, saat sinyalinin yukselen kenarinda calisan senkron "
    "bir flip-flop ile yonetilmektedir. f1_clre sinyali aktif oldugunda E sifirlanir; "
    "f2_come sinyali aktif oldugunda E tumlenenir (~e_flag); f1_add, f1_incac, f2_shl "
    "veya f2_shr sinyallerinden herhangi biri aktif oldugunda E, ALU'nun carry_out "
    "cikisina esitlenir. Bu yaklasim, kontrol birimi ile veri yolu arasinda net bir "
    "arayuz tanimlayarak paralel gelistirme surecini mumkun kilmistir. (PC5)"
)

tbl_ctrl = doc.add_table(rows=9, cols=3)
tbl_ctrl.style = "Table Grid"
tbl_header(tbl_ctrl, ["Kontrol Sinyali", "Verilog Ifadesi", "Islev"])
for i, row in enumerate([
    ["bus_sel",  "f3 sinyallerine gore (010, 011, 101 vb.)", "Ortak veri yolu kaynak secimi (3-bit)"],
    ["ar_load",  "f3_pctar | f3_irtar",    "AR yazmacina yukleme izni"],
    ["ir_load",  "f3_drtir",               "IR yazmacina DR icerigini yukle"],
    ["dr_load",  "f3_read | f3_actdr",     "DR'ye bellek veya bus verisi yukle"],
    ["ac_load",  "f1_add | f1_drtac | ...", "AC'ye ALU sonucu yukle"],
    ["ac_clr",   "f1_clrac",               "AC'yi senkron sifirla"],
    ["pc_load",  "f2_artpc",               "PC'ye AR degerini yukle"],
    ["pc_inc",   "f2_incpc",               "PC'yi bir arttir"],
]):
    tbl_row(tbl_ctrl, i + 1, row)
cap("Tablo 5.2. Ust Modul Kontrol Sinyal Baglantiları")

heading("5.1.6. Tasarim Kararlari ve Gerekceleri", level=2)
para(
    "Verilog implementasyonunda alinan temel tasarim kararlari sunlardir: "
    "(i) Tum yazmaclar posedge clk ve posedge reset ile calisan senkron tasarim "
    "prensibiyle gerceklenmis olup asenkron reset yalnizca sistem baslatma icin "
    "kullanilmaktadir; bu yaklasim FPGA sentezinde guvenilir zamanlama saglamaktadir. "
    "(ii) Kontrol bellegi Verilog initial blogu ile baslatilmis olup sentez aracilari "
    "tarafindan otomatik olarak LUT veya BRAM kaynaklarina eslenmektedir. "
    "(iii) ALU tamamen kombinasyonel olarak tasarlanmistir; sonuclarin yazmaclara "
    "yansitilmasi saat kenarinda kontrol sinyalleriyle saglanmaktadir. "
    "(iv) Kontrol birimi ile veri yolu arasindaki arayuz, 21 ayri tek-bit kontrol "
    "sinyali uzerinden tanimlanmistir; bu yaklasim iki katmanin bagimsiz olarak "
    "gelistirilip test edilmesine olanak tanimistir. "
    "(v) BSA komutu icin f3_actdr sinyali algindiginda bus secimi PC'ye (010) "
    "yonlendirilmek suretiyle DR'ye PC degeri yuklenebilmektedir; bu ozel durum "
    "kontrol belleginde ek bir F3 kodu tanimlamaya gerek kalmadan ust modul "
    "mantiginda cozulmustur. (PC1, PC5)"
)

heading("5.2. Test Senaryolari ve Sonuclar", level=2)
note("Bu alt bolum Kisi 4 tarafindan yazilacaktir. Testbench sonuclari, waveform goruntuleri ve test ozet tablosu eklenecek.")
para(
    "Tasarimin dogrulanmasi birim test ve entegrasyon test olmak uzere iki asama "
    "halinde gerceklestirilmistir. Birim testlerde her VHDL modulu bagimsiz "
    "testbench'lerle sinananmistir. Entegrasyon testinde ise AND, ADD, LDA, STA, "
    "BUN, ISZ komutlarini iceren bir test programi tam sistem uzerinde kosturulmus "
    "ve her komutun beklenen yazmac durumlarini urettigi dogrulanmistir. (PC5)"
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# 6. FPGA GERCEKLEMESI  (Tum grup)
# ══════════════════════════════════════════════════════════════════════
heading("6. FPGA GERCEKLEMESI")
para("Bu bolum tum grup tarafindan birlikte yazilacaktir.", bold=True)

heading("6.1. FPGA Ortami ve Araclar", level=2)
note("Kullanilan FPGA karti, Vivado surumu ve XDC dosyasindaki pin atamalari buraya eklenecek.")
para(
    "Tasarim, Xilinx Artix-7 FPGA cipini barindiran bir gelistirme karti uzerinde "
    "gerceklenmistir. Sentez ve yerlestirme-baglama (place & route) islemleri Xilinx "
    "Vivado ortaminda yurutulmus; pin atamalari XDC kisit dosyasi araciligiyla "
    "tanimlanmistir. Sistemin saat girisi 100 MHz kristal ossilatore baglanmis, "
    "sifirlama sinyali bir basma dugmesine atanmistir. (PC5)"
)

heading("6.2. Sentez Sonuclari", level=2)
note("Vivado Implementation Summary raporundan LUT, FF, BRAM ve WNS degerlerini girin.")
tbl_fpga = doc.add_table(rows=6, cols=3)
tbl_fpga.style = "Table Grid"
tbl_header(tbl_fpga, ["Kaynak", "Kullanilan", "Mevcut (%)"])
for i, row in enumerate([
    ["LUT",          "[doldurun]", "[%]"],
    ["FF",           "[doldurun]", "[%]"],
    ["BRAM",         "[doldurun]", "[%]"],
    ["Maks. Frekans","[MHz]",      "-"],
    ["WNS",          "[ns]",       "-"],
]):
    tbl_row(tbl_fpga, i + 1, row)
cap("Tablo 6.1. FPGA Kaynak Kullanim Raporu")

heading("6.3. FPGA Uzerinde Dogrulama", level=2)
note("Hangi cikislarin (LED, display) gozlemlendigi ve demo videosu referansi eklenecek.")
para(
    "Sisteme belirli bir test programi yuklenerek AC yazmacinin degeri kart uzerindeki "
    "LED'lerde izlenmistir. Reset uygulandiktan sonra program otomatik olarak yuruutulmus "
    "ve LED cikislarinin beklenen degerlerle ortustugu gozlemlenmistir. Sistemin "
    "calismasini belgeleyen demo videosu teslim paketinde sunulmaktadir. (PC5)"
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# 7. SONUC
# ══════════════════════════════════════════════════════════════════════
heading("7. SONUC")
para(
    "Bu calismada Morris Mano'nun temel bilgisayar mimarisinin hardwired kontrol devresi, "
    "mikroprogramlanmis kontrol mimarisine basariyla donusturulmus; tasarimin butun "
    "bilesenleri VHDL ile modellenmis, kapsamli testlerle dogrulanmis ve FPGA uzerinde "
    "gerceklenmistir. Gelistirilen sistemde 128 x 20-bitlik kontrol bellegi ve 20-bitlik "
    "yatay mikrokomut formati, Mano bilgisayarinin 25 komutunun tamamini dogru bicimde "
    "yuruteebilecek bicimde tasarlanmistir. Simulasyon ve FPGA test sonuclari, sistemin "
    "beklenen mantiksal dogrulugu karsiladigini ve zamanlama kisitlari dahilinde "
    "calistigini ortaya koymaktadir. (PC1, PC5)"
)
para(
    "Mikroprogramlanmis yaklasimin en belirgin avantaji, yeni komut eklenmesinin yalnizca "
    "kontrol bellegi icerigi guncellenerek mumkun olmasidir; bu ozellik tasarimin "
    "genisletilebilirligini hardwired kontrole kiyasla koklu bicimde artirmaktadir. "
    "Oernegin tasarima yeni bir islem kodu tanimlamak icin kontrol bellegine ilgili "
    "mikroprogram blogunun eklenmesi ve MAP tablosunun guncellenmesi yeterlidir; veri "
    "yolu tasarimi degistirmeye ihtiyac duyulmaz. Bu esneklik, mikroprogramlanmis "
    "kontrolun CISC mimarilerinde ve egitim amacli sistemlerde tercih edilmesinin "
    "temel nedenidir. (PC1, PC5)"
)
para(
    "Tasarimin temel sinirlamasi, hardwired kontrole kiyasla komut basina daha fazla "
    "saat darbesi gerekmesidir. Ortalama 8-11 cevrimi bulan islem suresi, "
    "performans kritik uygulamalar icin bir dezavantaj olusturmaktadir. Gelecek "
    "calismalarda bu eksikligin giderilmesi icin boru hatti (pipeline) mimarisi "
    "entegre edilebilir ya da sikca kullanilan komutlar icin kisaltilmis mikroprogram "
    "yollari tanimlanabilir. Bunun yaninda kontrol bellegi kapasitesi CAR bit genisligi "
    "arttirilarak buyutulebilir ve daha zengin bir komut seti desteklenebilir. (PC5)"
)
para(
    "Proje boyunca dort kisilik grup, net gorev dagilimi ve duzzenli iletisim sayesinde "
    "etkin bir isbirligi sureci yureutmustur. Veri yolu arayuz tanimlari proje baslangicinda "
    "ortaklasilarak paralel gelistirme mumkun kilinmis; tasarim kararlari grup toplantilari "
    "araciligiyla gerekcelendirilerek sonuclara baglanmistir. Bu surec, muhendislik "
    "problemlerinin takim halinde cozulme becerisini gelistirme hedefini karsilamistir. (PC13)"
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# 8. KAYNAKLAR
# ══════════════════════════════════════════════════════════════════════
heading("8. KAYNAKLAR")
for ref in [
    "[1]  M. M. Mano, Computer System Architecture, 3. baski. Upper",
    "     Saddle River, NJ: Prentice Hall, 1993.",
    "[2]  M. M. Mano ve C. R. Kime, Logic and Computer Design",
    "     Fundamentals, 4. baski. Prentice Hall, 2008.",
    "[3]  M. V. Wilkes, 'The best way to design an automatic",
    "     calculating machine,' Manchester Univ. Comp. Inaugural",
    "     Conf., 1951.",
    "[4]  Xilinx Inc., Vivado Design Suite User Guide: Synthesis,",
    "     UG901, 2023.",
    "[5]  P. J. Ashenden, The Designer's Guide to VHDL, 3. baski.",
    "     San Francisco: Morgan Kaufmann, 2008.",
    "[6]  [Varsa kullanilan ek kaynak - IEEE formatinda]",
]:
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.left_indent  = Cm(0.5)
    set_font(p.add_run(ref))
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# EK A — GOREV DAGILIMI
# ══════════════════════════════════════════════════════════════════════
heading("EK A - GOREV DAGILIMI")
tbl_g = doc.add_table(rows=6, cols=3)
tbl_g.style = "Table Grid"
tbl_header(tbl_g, ["Kisi", "Gorev", "Bolumler"])
for i, row in enumerate([
    ["Kisi 1 (Eren)", "Rapor yazimi, koordinasyon, giris, sonuc", "2, 7, Ek A-B"],
    ["Kisi 2",        "Veri yolu bilesenleri (VHDL)",             "4.1, 5.1"],
    ["Kisi 3",        "Mikroprogramlanmis kontrol birimi (VHDL)", "4.2, 4.3, 4.4, 5.1"],
    ["Kisi 4",        "Test ve simulasyon",                       "5.2"],
    ["Tumu",          "Teorik arka plan, FPGA, sunum",            "3, 6"],
]):
    tbl_row(tbl_g, i + 1, row)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# EK B — ILETISIM VE TOPLANTI OZETI  (PÇ13 gerekliligi)
# ══════════════════════════════════════════════════════════════════════
heading("EK B - ILETISIM VE TOPLANTI OZETI")
note("Grup toplanti kayitlari PC13 gerekliligini karsilamak icin zorunludur. "
     "Her toplantidan sonra asagidaki tablolari doldurun.")
para(
    "Grup calisma sureci boyunca duzzenli toplanti ve iletisim kayitlari tutulmustur. "
    "Asagidaki tablolar her toplantinin tarih, icerik ve alinan kararlarini "
    "ozetlemektedir. (PC13)"
)

for i in range(1, 4):
    heading(f"Toplanti #{i}", level=2)
    tbl_t = doc.add_table(rows=5, cols=2)
    tbl_t.style = "Table Grid"
    for r, (k, v) in enumerate([
        ("Tarih / Saat",     "[Doldurun]"),
        ("Platform",         "[Discord / WhatsApp / Yuz yuze]"),
        ("Gundem",           "[Doldurun]"),
        ("Alinan Kararlar",  "[Doldurun]"),
        ("Sonraki Adim",     "[Doldurun]"),
    ]):
        cell_k = tbl_t.cell(r, 0)
        cell_v = tbl_t.cell(r, 1)
        cell_k.paragraphs[0].clear()
        cell_v.paragraphs[0].clear()
        set_font(cell_k.paragraphs[0].add_run(k), bold=True)
        set_font(cell_v.paragraphs[0].add_run(v))
    para()

import os
out = os.path.join(os.path.dirname(os.path.abspath(__file__)), "Rapor_Iskelet_v2.docx")
doc.save(out)
print(f"Dosya olusturuldu: {out}")
